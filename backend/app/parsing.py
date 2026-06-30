import os
import io
import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from app.config import ALLOWED_DOC_EXTENSIONS, ALLOWED_CERT_EXTENSIONS, MAX_UPLOAD_MB


class UnsupportedFileError(Exception):
    """Raised so the API layer can return a clear error (FR-1.7)."""


def validate_extension(filename: str, doc_type: str):
    ext = os.path.splitext(filename.lower())[1]
    allowed = ALLOWED_CERT_EXTENSIONS if doc_type == "certificate" else ALLOWED_DOC_EXTENSIONS
    if ext not in allowed:
        raise UnsupportedFileError(
            f"Unsupported file type '{ext}' for {doc_type}. Allowed: {', '.join(sorted(allowed))}"
        )
    return ext


def validate_size(file_bytes: bytes):
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_UPLOAD_MB:
        raise UnsupportedFileError(f"File exceeds the {MAX_UPLOAD_MB}MB limit (FR-1.1).")


def parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def parse_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def parse_image_ocr(file_bytes: bytes) -> str:
    """FR-1.4: OCR certificates to extract issuer/title/date as free text."""
    img = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(img).strip()


def parse_uploaded_file(filename: str, doc_type: str, file_bytes: bytes) -> str:
    validate_size(file_bytes)
    ext = validate_extension(filename, doc_type)
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    if ext == ".docx":
        return parse_docx(file_bytes)
    if ext in (".jpg", ".jpeg", ".png"):
        return parse_image_ocr(file_bytes)
    raise UnsupportedFileError(f"No parser available for '{ext}'.")


def fetch_and_summarize_url(url: str, max_chars: int = 6000) -> str:
    """FR-1.3: fetch a public project/portfolio link and extract readable text.
    Summarization into the knowledge base proper happens later via the LLM
    at clone-generation time; here we just extract clean text."""
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "DigitalCloneBot/1.0"})
        resp.raise_for_status()
    except requests.RequestException as e:
        raise UnsupportedFileError(f"Could not fetch URL '{url}': {e}")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    text = " ".join(soup.get_text(separator=" ").split())
    return text[:max_chars]
